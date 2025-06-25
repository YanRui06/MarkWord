#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
跨平台测试脚本 - 验证Markdown转Word转换器的基本功能
支持 Windows, macOS, Linux
"""

import sys
import os
import platform
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'src'))

def test_platform_detection():
    """测试平台检测功能"""
    print("检测平台信息...")
    system = platform.system()
    print(f"  操作系统: {system}")
    print(f"  Python版本: {platform.python_version()}")
    print(f"  架构: {platform.machine()}")
    print(f"  平台: {platform.platform()}")
    return True

def test_basic_conversion():
    """测试基本转换功能"""
    print("开始测试Markdown转Word转换...")
    
    try:
        # 导入必要的模块
        import markdown
        from docx import Document
        from bs4 import BeautifulSoup
        print("✅ 所有依赖模块导入成功")
        
        # 测试Markdown解析
        md_text = """# 测试标题
这是一个测试段落。

## 二级标题
- 列表项1
- 列表项2

**粗体文本** 和 *斜体文本*

```python
def hello():
    print("Hello World")
```
"""
        
        # 转换为HTML
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'toc', 'tables'])
        html = md.convert(md_text)
        print("✅ Markdown解析成功")
        
        # 创建Word文档
        doc = Document()
        doc.add_heading("跨平台测试文档", 0)
        doc.add_paragraph("这是一个跨平台测试段落")
        
        # 根据系统设置字体
        system = platform.system()
        if system == "Windows":
            font_name = '微软雅黑'
        elif system == "Darwin":  # macOS
            font_name = 'PingFang SC'
        else:  # Linux
            font_name = 'DejaVu Sans'
            
        # 添加系统信息
        info_para = doc.add_paragraph(f"测试系统: {platform.platform()}")
        
        # 保存测试文档
        test_output = f"test_output_{system.lower()}.docx"
        doc.save(test_output)
        print(f"✅ Word文档创建成功: {test_output}")
        
        # 检查文件是否存在
        if os.path.exists(test_output):
            file_size = os.path.getsize(test_output)
            print(f"✅ 文件大小: {file_size} 字节")
            print("🎉 基本功能测试通过!")
            return True
        else:
            print("❌ 文件创建失败")
            return False
            
    except Exception as e:
        print(f"❌ 测试失败: {str(e)}")
        return False

def test_gui_import():
    """测试GUI相关模块导入"""
    try:
        import tkinter as tk
        print("✅ Tkinter导入成功")
        
        # 测试创建根窗口
        root = tk.Tk()
        root.withdraw()  # 隐藏窗口
        
        # 测试字体设置
        system = platform.system()
        if system == "Windows":
            test_font = ("微软雅黑", 10)
        elif system == "Darwin":
            test_font = ("SF Pro Display", 10)
        else:
            test_font = ("DejaVu Sans", 10)
            
        try:
            test_label = tk.Label(root, text="Test", font=test_font)
            print(f"✅ 字体测试成功: {test_font[0]}")
        except:
            print(f"⚠️  字体 {test_font[0]} 不可用，将使用默认字体")
        
        print("✅ Tkinter窗口创建成功")
        root.destroy()
        return True
    except Exception as e:
        print(f"❌ GUI测试失败: {str(e)}")
        if platform.system() == "Linux":
            print("提示: Linux用户可能需要安装python3-tk:")
            print("  Ubuntu/Debian: sudo apt-get install python3-tk")
            print("  CentOS/RHEL: sudo yum install tkinter")
            print("  Fedora: sudo dnf install python3-tkinter")
        return False

def test_file_operations():
    """测试文件操作"""
    try:
        from pathlib import Path
        
        # 测试文件路径处理
        test_path = Path("test_file.txt")
        test_path.write_text("测试内容", encoding='utf-8')
        
        if test_path.exists():
            content = test_path.read_text(encoding='utf-8')
            test_path.unlink()  # 删除测试文件
            print("✅ 文件操作测试成功")
            return True
        return False
    except Exception as e:
        print(f"❌ 文件操作测试失败: {str(e)}")
        return False

if __name__ == "__main__":
    print("=" * 60)
    print("Markdown转Word转换器 - 跨平台功能测试")
    print("=" * 60)
    
    # 测试平台检测
    platform_test = test_platform_detection()
    print()
    
    # 测试基本转换功能
    basic_test = test_basic_conversion()
    print()
    
    # 测试GUI功能
    gui_test = test_gui_import()
    print()
    
    # 测试文件操作
    file_test = test_file_operations()
    print()
    
    # 总结
    all_tests = [platform_test, basic_test, gui_test, file_test]
    passed = sum(all_tests)
    total = len(all_tests)
    
    print("=" * 60)
    print(f"测试结果: {passed}/{total} 项测试通过")
    
    if all(all_tests):
        print("🎉 所有测试通过! 可以启动完整程序。")
        print()
        print("启动方式:")
        print("1. 使用启动器: python launcher.py")
        print("2. 直接运行: python md_to_word_converter.py")
        if platform.system() == "Windows":
            print("3. Windows用户: 双击 start.bat")
        else:
            print("3. macOS/Linux用户: ./start.sh")
    else:
        print("❌ 部分测试失败，请检查环境配置。")
        
    print("=" * 60)
