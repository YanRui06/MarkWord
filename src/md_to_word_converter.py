"""
Markdown到Word转换器
支持保持原有格式的Markdown文档转换
跨平台支持：Windows、macOS、Linux
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import sys
import platform
import threading
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import re


class MarkdownToWordConverter:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Markdown转Word转换器")
        self.root.geometry("800x600")
        
        # 检测操作系统并设置相应的配置
        self.system = platform.system()
        self.setup_platform_config()
        
        # 设置窗口图标（如果存在）
        self.setup_window_icon()
        
        # 设置样式
        self.setup_styles()
        
        self.setup_ui()
        
    def setup_platform_config(self):
        """根据操作系统设置平台相关配置"""
        if self.system == "Windows":
            self.default_font = ("微软雅黑", 10)
            self.title_font = ("微软雅黑", 20, "bold")
            self.code_font = ("Consolas", 9)
            self.bg_color = '#f0f0f0'
        elif self.system == "Darwin":  # macOS
            self.default_font = ("SF Pro Display", 10)
            self.title_font = ("SF Pro Display", 20, "bold")
            self.code_font = ("SF Mono", 9)
            self.bg_color = '#f5f5f5'
        else:  # Linux
            self.default_font = ("DejaVu Sans", 10)
            self.title_font = ("DejaVu Sans", 20, "bold")
            self.code_font = ("DejaVu Sans Mono", 9)
            self.bg_color = '#f0f0f0'
            
        # 设置窗口背景色
        self.root.configure(bg=self.bg_color)
        
    def setup_window_icon(self):
        """设置窗口图标"""
        try:
            # 尝试设置窗口图标（如果有的话）
            if self.system == "Windows":
                # Windows可以使用.ico文件
                icon_path = Path(__file__).parent / "icon.ico"
                if icon_path.exists():
                    self.root.iconbitmap(str(icon_path))
            else:
                # macOS和Linux使用PhotoImage
                icon_path = Path(__file__).parent / "icon.png"
                if icon_path.exists():
                    icon = tk.PhotoImage(file=str(icon_path))
                    self.root.iconphoto(False, icon)
        except Exception:
            pass  # 忽略图标设置错误
            
    def setup_styles(self):
        """设置样式"""
        self.style = ttk.Style()
        
        # 根据系统选择合适的主题
        if self.system == "Windows":
            try:
                self.style.theme_use('vista')
            except:
                self.style.theme_use('clam')
        elif self.system == "Darwin":
            try:
                self.style.theme_use('aqua')
            except:
                self.style.theme_use('clam')
        else:  # Linux
            try:
                self.style.theme_use('clam')
            except:
                self.style.theme_use('default')
        
    def setup_ui(self):
        """设置用户界面"""
        # 主标题
        title_label = tk.Label(
            self.root, 
            text="Markdown转Word转换器", 
            font=self.title_font,
            bg=self.bg_color,
            fg='#2c3e50'
        )
        title_label.pack(pady=20)
        
        # 文件选择框架
        file_frame = tk.Frame(self.root, bg=self.bg_color)
        file_frame.pack(pady=10, padx=20, fill='x')
        
        # 输入文件选择
        input_frame = tk.Frame(file_frame, bg=self.bg_color)
        input_frame.pack(fill='x', pady=5)
        
        tk.Label(input_frame, text="选择Markdown文件:", font=self.default_font, bg=self.bg_color).pack(anchor='w')
        
        input_path_frame = tk.Frame(input_frame, bg=self.bg_color)
        input_path_frame.pack(fill='x', pady=5)
        
        self.input_path_var = tk.StringVar()
        self.input_entry = tk.Entry(
            input_path_frame, 
            textvariable=self.input_path_var, 
            font=self.default_font,
            width=60
        )
        self.input_entry.pack(side='left', fill='x', expand=True)
        
        input_browse_btn = ttk.Button(
            input_path_frame, 
            text="浏览", 
            command=self.browse_input_file
        )
        input_browse_btn.pack(side='right', padx=(5, 0))
        
        # 输出文件选择
        output_frame = tk.Frame(file_frame, bg=self.bg_color)
        output_frame.pack(fill='x', pady=5)
        
        tk.Label(output_frame, text="保存Word文件为:", font=self.default_font, bg=self.bg_color).pack(anchor='w')
        
        output_path_frame = tk.Frame(output_frame, bg=self.bg_color)
        output_path_frame.pack(fill='x', pady=5)
        
        self.output_path_var = tk.StringVar()
        self.output_entry = tk.Entry(
            output_path_frame, 
            textvariable=self.output_path_var, 
            font=self.default_font,
            width=60
        )
        self.output_entry.pack(side='left', fill='x', expand=True)
        
        output_browse_btn = ttk.Button(
            output_path_frame, 
            text="浏览", 
            command=self.browse_output_file
        )
        output_browse_btn.pack(side='right', padx=(5, 0))
        
        # 选项框架
        options_frame = tk.LabelFrame(
            self.root, 
            text="转换选项", 
            font=(self.default_font[0], 12, "bold"),
            bg=self.bg_color,
            fg='#2c3e50'
        )
        options_frame.pack(pady=20, padx=20, fill='x')
        
        # 选项复选框
        self.preserve_formatting = tk.BooleanVar(value=True)
        tk.Checkbutton(
            options_frame, 
            text="保持原有格式", 
            variable=self.preserve_formatting,
            font=self.default_font,
            bg=self.bg_color
        ).pack(anchor='w', padx=10, pady=5)
        
        self.include_toc = tk.BooleanVar(value=False)
        tk.Checkbutton(
            options_frame, 
            text="生成目录", 
            variable=self.include_toc,
            font=self.default_font,
            bg=self.bg_color
        ).pack(anchor='w', padx=10, pady=5)
        
        self.process_images = tk.BooleanVar(value=True)
        tk.Checkbutton(
            options_frame, 
            text="处理图片", 
            variable=self.process_images,
            font=self.default_font,
            bg=self.bg_color
        ).pack(anchor='w', padx=10, pady=5)
        
        self.clean_formatting = tk.BooleanVar(value=True)
        tk.Checkbutton(
            options_frame, 
            text="清理格式", 
            variable=self.clean_formatting,
            font=self.default_font,
            bg=self.bg_color
        ).pack(anchor='w', padx=10, pady=5)
        
        # 转换按钮
        convert_btn = ttk.Button(
            self.root, 
            text="开始转换", 
            command=self.start_conversion
        )
        convert_btn.pack(pady=20)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.root, 
            variable=self.progress_var, 
            maximum=100,
            length=400
        )
        self.progress_bar.pack(pady=10)
        
        # 状态标签
        self.status_var = tk.StringVar(value="就绪")
        self.status_label = tk.Label(
            self.root, 
            textvariable=self.status_var, 
            font=self.default_font,
            bg=self.bg_color,
            fg='#27ae60'
        )
        self.status_label.pack(pady=5)
        
        # 日志文本框
        log_frame = tk.LabelFrame(
            self.root, 
            text="转换日志", 
            font=self.default_font,
            bg=self.bg_color
        )
        log_frame.pack(pady=10, padx=20, fill='both', expand=True)
        
        self.log_text = tk.Text(
            log_frame, 
            height=8, 
            font=self.code_font,
            bg='#2c3e50',
            fg='#ecf0f1',
            insertbackground='#ecf0f1'
        )
        scrollbar = tk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
    def browse_input_file(self):
        """浏览输入文件"""
        filename = filedialog.askopenfilename(
            title="选择Markdown文件",
            filetypes=[("Markdown文件", "*.md *.markdown"), ("所有文件", "*.*")]
        )
        if filename:
            self.input_path_var.set(filename)
            # 自动设置输出文件名
            if not self.output_path_var.get():
                output_name = Path(filename).stem + ".docx"
                output_path = Path(filename).parent / output_name
                self.output_path_var.set(str(output_path))
                
    def browse_output_file(self):
        """浏览输出文件"""
        filename = filedialog.asksaveasfilename(
            title="保存Word文件",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.output_path_var.set(filename)
            
    def log_message(self, message):
        """添加日志消息"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        
    def update_status(self, message):
        """更新状态"""
        self.status_var.set(message)
        self.root.update_idletasks()
        
    def update_progress(self, value):
        """更新进度条"""
        self.progress_var.set(value)
        self.root.update_idletasks()
        
    def start_conversion(self):
        """开始转换"""
        input_file = self.input_path_var.get().strip()
        output_file = self.output_path_var.get().strip()
        
        if not input_file:
            messagebox.showerror("错误", "请选择输入的Markdown文件")
            return
            
        if not output_file:
            messagebox.showerror("错误", "请指定输出的Word文件")
            return
            
        if not os.path.exists(input_file):
            messagebox.showerror("错误", "输入文件不存在")
            return
            
        # 在新线程中执行转换
        self.log_text.delete(1.0, tk.END)
        threading.Thread(target=self.convert_file, args=(input_file, output_file), daemon=True).start()
        
    def convert_file(self, input_file, output_file):
        """转换文件"""
        try:
            self.update_status("正在转换...")
            self.update_progress(0)
            
            # 读取Markdown文件
            self.log_message("正在读取Markdown文件...")
            with open(input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            self.update_progress(20)
            
            # 转换Markdown为HTML
            self.log_message("正在解析Markdown...")
            md = markdown.Markdown(extensions=[
                'extra', 'codehilite', 'toc', 'tables', 'fenced_code'
            ])
            html_content = md.convert(md_content)
            self.update_progress(40)
            
            # 创建Word文档
            self.log_message("正在创建Word文档...")
            doc = Document()
            self.setup_document_styles(doc)
            self.update_progress(50)
            
            # 解析HTML并转换为Word
            self.log_message("正在转换内容...")
            soup = BeautifulSoup(html_content, 'html.parser')
            self.html_to_docx(soup, doc, Path(input_file).parent)
            self.update_progress(80)
            
            # 保存文档
            self.log_message(f"正在保存到: {output_file}")
            doc.save(output_file)
            self.update_progress(100)
            
            self.update_status("转换完成")
            self.log_message("转换成功完成!")
            messagebox.showinfo("成功", f"转换完成!\n文件已保存为: {output_file}")
            
        except Exception as e:
            self.update_status("转换失败")
            self.log_message(f"错误: {str(e)}")
            messagebox.showerror("错误", f"转换失败: {str(e)}")
            
    def setup_document_styles(self, doc):
        """设置文档样式"""
        styles = doc.styles
        
        # 根据操作系统选择合适的字体
        if self.system == "Windows":
            font_name = '微软雅黑'
            code_font_name = 'Consolas'
        elif self.system == "Darwin":  # macOS
            font_name = 'PingFang SC'
            code_font_name = 'SF Mono'
        else:  # Linux
            font_name = 'DejaVu Sans'
            code_font_name = 'DejaVu Sans Mono'
        
        # 设置标题样式
        for i in range(1, 7):
            try:
                heading_style = styles[f'Heading {i}']
                heading_style.font.name = font_name
                heading_style.font.size = Pt(24 - i * 2)
                heading_style.font.bold = True
                heading_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            except Exception as e:
                self.log_message(f"设置标题样式时警告: {str(e)}")
            
        # 设置正文样式
        try:
            normal_style = styles['Normal']
            normal_style.font.name = font_name
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15
        except Exception as e:
            self.log_message(f"设置正文样式时警告: {str(e)}")
            
        # 设置代码字体（用于后续代码块）
        self.document_code_font = code_font_name
        
    def html_to_docx(self, soup, doc, base_path):
        """将HTML转换为Word文档"""
        for element in soup.children:
            if hasattr(element, 'name'):
                self.process_element(element, doc, base_path)
                
    def process_element(self, element, doc, base_path):
        """处理HTML元素"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # 处理标题
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text().strip(), level=level)
            
        elif element.name == 'p':
            # 处理段落
            paragraph = doc.add_paragraph()
            self.process_inline_elements(element, paragraph)
            
        elif element.name == 'ul':
            # 处理无序列表
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(style='List Bullet')
                self.process_inline_elements(li, paragraph)
                
        elif element.name == 'ol':
            # 处理有序列表
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(style='List Number')
                self.process_inline_elements(li, paragraph)
                
        elif element.name == 'blockquote':
            # 处理引用
            paragraph = doc.add_paragraph()
            paragraph.style = 'Quote'
            self.process_inline_elements(element, paragraph)
            
        elif element.name == 'pre':
            # 处理代码块
            code_text = element.get_text()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code_text)
            run.font.name = getattr(self, 'document_code_font', 'Courier New')
            run.font.size = Pt(10)
            paragraph.style = 'No Spacing'
            
        elif element.name == 'table':
            # 处理表格
            self.process_table(element, doc)
            
        elif element.name == 'img' and self.process_images.get():
            # 处理图片
            self.process_image(element, doc, base_path)
            
        elif element.name == 'hr':
            # 处理分隔线
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.LINE)
            
    def process_inline_elements(self, element, paragraph):
        """处理内联元素"""
        for child in element.children:
            if hasattr(child, 'name'):
                if child.name == 'strong' or child.name == 'b':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = getattr(self, 'document_code_font', 'Courier New')
                    run.font.size = Pt(10)
                elif child.name == 'a':
                    # 处理链接
                    text = child.get_text()
                    url = child.get('href', '')
                    if url:
                        run = paragraph.add_run(f"{text} ({url})")
                        run.font.color.rgb = RGBColor(0x00, 0x7a, 0xcc)
                    else:
                        run = paragraph.add_run(text)
                else:
                    paragraph.add_run(child.get_text())
            else:
                # 纯文本
                paragraph.add_run(str(child))
                
    def process_table(self, table_element, doc):
        """处理表格"""
        rows = table_element.find_all('tr')
        if not rows:
            return
            
        # 计算列数
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        # 创建表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table.cell(i, j).text = cell.get_text().strip()
                    # 如果是表头，设置粗体
                    if cell.name == 'th':
                        for paragraph in table.cell(i, j).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                
    def process_image(self, img_element, doc, base_path):
        """处理图片"""
        try:
            src = img_element.get('src', '')
            alt = img_element.get('alt', '')
            
            if src:
                # 处理相对路径
                if not src.startswith(('http://', 'https://')):
                    img_path = base_path / src
                    if img_path.exists():
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(str(img_path), width=Inches(6))
                        
                        # 添加图片说明
                        if alt:
                            caption = doc.add_paragraph(f"图片: {alt}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
        except Exception as e:
            self.log_message(f"处理图片时出错: {str(e)}")
            
    def run(self):
        """运行应用程序"""
        self.root.mainloop()


if __name__ == "__main__":
    app = MarkdownToWordConverter()
    app.run()
